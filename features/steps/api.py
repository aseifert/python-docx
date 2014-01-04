# encoding: utf-8

"""
Step implementations for basic API features
"""

from behave import then, when

from docx.shared import Inches

from .helpers import test_file_path


# when ====================================================

@when('I add a heading specifying level={level_str}')
def when_add_heading_specifying_level(context, level_str):
    level = int(level_str)
    document = context.document
    document.add_heading(level=level)


@when('I add a heading specifying only its text')
def when_add_heading_specifying_only_its_text(context):
    document = context.document
    context.heading_text = 'Spam vs. Eggs'
    document.add_heading(context.heading_text)


@when('I add a page break to the document')
def when_add_page_break_to_document(context):
    document = context.document
    document.add_page_break()


@when('I add a paragraph specifying its style')
def when_add_paragraph_specifying_style(context):
    document = context.document
    context.paragraph_style = 'barfoo'
    document.add_paragraph(style=context.paragraph_style)


@when('I add a paragraph specifying its text')
def when_add_paragraph_specifying_text(context):
    document = context.document
    context.paragraph_text = 'foobar'
    document.add_paragraph(context.paragraph_text)


@when('I add a paragraph without specifying text or style')
def when_add_paragraph_without_specifying_text_or_style(context):
    document = context.document
    document.add_paragraph()


@when('I add a picture specifying 1.75" width and 2.5" height')
def when_add_picture_specifying_width_and_height(context):
    document = context.document
    context.picture = document.add_picture(
        test_file_path('monty-truth.png'),
        width=Inches(1.75), height=Inches(2.5)
    )


@when('I add a picture specifying a height of 1.5 inches')
def when_add_picture_specifying_height(context):
    document = context.document
    context.picture = document.add_picture(
        test_file_path('monty-truth.png'), height=Inches(1.5)
    )


@when('I add a picture specifying a width of 1.5 inches')
def when_add_picture_specifying_width(context):
    document = context.document
    context.picture = document.add_picture(
        test_file_path('monty-truth.png'), width=Inches(1.5)
    )


@when('I add a picture specifying only the image file')
def when_add_picture_specifying_only_image_file(context):
    document = context.document
    context.picture = document.add_picture(test_file_path('monty-truth.png'))


# then =====================================================

@then('the last paragraph contains only a page break')
def then_last_paragraph_contains_only_a_page_break(context):
    document = context.document
    p = document.paragraphs[-1]
    assert len(p.runs) == 1
    assert len(p.runs[0]) == 1
    assert p.runs[0]._r[0].type == 'page'


@then('the last paragraph contains the heading text')
def then_last_p_contains_heading_text(context):
    document = context.document
    text = context.heading_text
    p = document.paragraphs[-1]
    assert p.text == text


@then('the last paragraph contains the text I specified')
def then_last_p_contains_specified_text(context):
    document = context.document
    text = context.paragraph_text
    p = document.paragraphs[-1]
    assert p.text == text


@then('the last paragraph has the style I specified')
def then_last_p_has_specified_style(context):
    document = context.document
    style = context.paragraph_style
    p = document.paragraphs[-1]
    assert p.style == style


@then('the last paragraph is the empty paragraph I added')
def then_last_p_is_empty_paragraph_added(context):
    document = context.document
    p = document.paragraphs[-1]
    assert p.text == ''


@then('the picture has its native width and height')
def then_picture_has_native_width_and_height(context):
    picture = context.picture
    assert picture.width == 1905000, 'got %d' % picture.width
    assert picture.height == 2717800, 'got %d' % picture.height


@then('the picture height is 2.14 inches')
def then_picture_height_is_value_2(context):
    picture = context.picture
    assert picture.height == 1956816, 'got %d' % picture.height


@then('the picture height is 2.5 inches')
def then_picture_height_is_value(context):
    picture = context.picture
    assert picture.height == 2286000, 'got %d' % picture.height


@then('the picture width is 1.05 inches')
def then_picture_width_is_value_2(context):
    picture = context.picture
    assert picture.width == 961402, 'got %d' % picture.width


@then('the picture width is 1.75 inches')
def then_picture_width_is_value(context):
    picture = context.picture
    assert picture.width == 1600200, 'got %d' % picture.width


@then('the style of the last paragraph is \'{style}\'')
def then_style_of_last_paragraph_is_style(context, style):
    document = context.document
    p = document.paragraphs[-1]
    assert p.style == style
